import os
import re
import sys
import json
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
import AcrossValidation


class AcrossReader:

    def __init__(self, segment_id, source, translation):
        self.id = segment_id
        self.source = source
        self.translation = translation

    @classmethod
    def read_htm_file(cls, htm_file):
        """
        Reads a file in htm format that is generated by Across and converts the text into a csv file that contains
        the segment id, the source text, and the translated text. "§" is used as a separator.
        """

        # Validation
        across_validator = AcrossValidation.AcrossValidation

        try:
            across_validator.check_file_ending(htm_file)
            across_validator.check_across_htm_file(htm_file)

            file_to_store = cls.__generate_destination_file__(htm_file)

            with open(htm_file, "r", encoding="utf-8") as file_to_be_read:
                lines = file_to_be_read.readlines()
                counter = 1
                translations = list()
                source = ''

                for entry in lines:
                    if entry.startswith("<DIV id=contents"):
                        tmp = entry

                        if counter == 2:
                            source = tmp[:-1]
                            source = cls.__remove_nbsp_tag__(source)
                            source = cls.__remove_tags_inline(source)
                            source = cls.__replace_special_characters__(source)
                            counter += 1

                        elif counter == 3:
                            translation = tmp
                            translation = cls.__remove_nbsp_tag__(translation)
                            translation = cls.__remove_tags_inline(translation)
                            translation = cls.__replace_special_characters__(translation)
                            counter = 1
                            translation_entry = AcrossReader(segment_id, source, translation)
                            translations.append(translation_entry)

                    if 'inactiveNumbering" width=' in entry and counter == 1:
                        segment_id = re.sub('<TD.*?MARGIN: 2px 0px 0px"><SPAN class=atom>', "", entry)
                        segment_id = re.sub("</SPAN></PRE></TD>", "", segment_id)[:-1]
                        counter += 1

                all_translations_to_store = cls.__remove_duplicates__(translations)

                cls.__write_file_to_docx__(all_translations_to_store, file_to_store)

        except FileExistsError as error:
            sys.exit(str(error))
        except BaseException as error:
            sys.exit((str(error)))

    @staticmethod
    def __set_col_widths__(table):
        """
        Helper method that sets the column width of a table of a docx file.
        :param table: table where the column width needs to be set.
        """

        widths = (Cm(1.3), Cm(7), Cm(7))
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

    @classmethod
    def __write_file_to_docx__(cls, translation_list_with_AcrossReader, document_name):
        """
        Takes a list containing elements of AcrossReader that contains of a segment id, a source text and a translation.
        Writes each entry of the given list in a table and saves the file as a docx document.

        :param translation_list_with_AcrossReader: list containing elements of AcrossReader
        :param document_name: name of the destination file
        """

        document = Document()

        # General Styling
        style = document.styles['Normal']
        font = style.font
        font.name = 'Browallia New'
        font.size = Pt(12)

        # Document title
        document.add_heading('Datei für zweisprachige Überprüfung')

        # Document content
        translation_tuples = list()
        for translation in translation_list_with_AcrossReader:
            translation_tuple = (translation.id, translation.source, translation.translation)
            translation_tuples.append(translation_tuple)

        table = document.add_table(rows=1, cols=3, style="Table Grid")
        table.autofit = False
        hdr_cells = table.rows[0].cells

        # Table header and layout
        header_list = list()
        header_list.append(hdr_cells[0].paragraphs[0].add_run('ID'))
        header_list.append(hdr_cells[1].paragraphs[0].add_run('Ausgangstext'))
        header_list.append(hdr_cells[2].paragraphs[0].add_run('Übersetzung'))

        for header in header_list:
            header.bold = True

        # Table content
        for segment_id, source, translation in translation_tuples:
            row_cells = table.add_row().cells
            row_cells[0].text = str(segment_id)
            row_cells[1].text = source
            row_cells[2].text = translation

        AcrossReader.__set_col_widths__(table)
        document.save(document_name)

    @classmethod
    def __remove_duplicates__(cls, translation_list):
        """
        Removes entries that occur multiple times in a given list containing translation entries (consist segment id,
        source text, and translation.

        :param translation_list: list containing all translations of the text. The entries are instances of the class
        AcrossReader.
        :return: list holding translations without duplicates.
        """

        list_without_duplicates = list()

        for entry in translation_list:
            exists = False
            for entry_without_duplicates in list_without_duplicates:
                if entry_without_duplicates.source == entry.source:
                    exists = True
                    break

            if exists is False:
                list_without_duplicates.append(entry)

        return list_without_duplicates

    @classmethod
    def __remove_tags_inline(cls, line):
        """
        Removes tags that occur within the text and removes them or replaces them with responding tags that are
        defined in a json file.

        :param line: one segment holding the source text or the translation.
        :return: line with customized tags (replaced or removed).
        """

        # default tags that are used in every htm file that is generated by Across
        line_cleaned = re.sub("<WBR>", "", line)
        line_cleaned = re.sub("<TR.*?>", "", line_cleaned)
        line_cleaned = re.sub("<TD.*?>", "", line_cleaned)
        line_cleaned = re.sub("<DIV.*?>", "", line_cleaned)
        line_cleaned = re.sub("<PRE.*?>", "", line_cleaned)
        line_cleaned = re.sub("<SPAN.*?>", "", line_cleaned)
        line_cleaned = re.sub("</SPAN>", "", line_cleaned)
        line_cleaned = re.sub("</PRE>", "", line_cleaned)
        line_cleaned = re.sub("</DIV>", "", line_cleaned)
        line_cleaned = re.sub("</TD>", "", line_cleaned)
        line_cleaned = re.sub("</TR>", "", line_cleaned)

        # customer specific tags of each project
        with open('tags.json', 'r', encoding="utf-8") as tag_file:
            tag_dict = json.load(tag_file)

        for tag in tag_dict.get('tags'):
            line_cleaned = cls.__replace_tag__(line_cleaned, tag.get('name'), tag.get('show_tag'), tag.get('opening'), tag.get('closing'))

        line_cleaned = re.sub('<IMG.*?alt="<Pfad>.*?</Pfad>.*?">', "<Pfad />", line_cleaned)
        line_cleaned = re.sub('<IMG.*?png">', "", line_cleaned)

        return line_cleaned

    @classmethod
    def __replace_tag__(cls, line, tag_name, show_tag, opening, closing):
        """
        Replaces all the tags with the given opening and closing tag as it is used in the source text.
        The tags that should be considered is stored in a json file that supports UTF 8 encoding. If the show_tag is
        set to false, the tag will be removed and not replaced by any defined name.

        :param line: line to be considered.
        :param tag_name: name that should appear in the cleaned file.
        :param show_tag: determines whether the tag should be replaced (true) or removed (false).
        :param opening: opening tag as it is used in the source document.
        :param closing: closing tag as it is used in the source document. If there is no closing tag, it needs to be
        set to an empty string "".
        :return: line that replaced / removed the tags according to the information in the json file.
        """

        counter = line.count(opening)
        counter_closing = line.count(closing)
        if counter == counter_closing:
            for occurrence in range(0, counter):
                if show_tag is True:
                    line = re.sub(f'<IMG.*?alt="{opening}".*?">', f"<{tag_name}>", line, 1)
                else:
                    line = re.sub(f'<IMG.*?alt="{opening}".*?">', '', line, 1)

                if "" != closing:
                    if show_tag is True:
                        line = re.sub(f'<IMG.*?alt="{closing}".*?">', f"</{tag_name}>", line, 1)
                    else:
                        line = re.sub(f'<IMG.*?alt="{closing}".*?">', '', line, 1)
        else:

            for occurrence in range(0, counter):
                if show_tag is True:
                    line = re.sub(f'<IMG.*?alt="{opening}".*?">', f"<{tag_name}>", line, 1)
                else:
                    line = re.sub(f'<IMG.*?alt="{opening}".*?">', '', line, 1)

            for occurrence in range(0, counter_closing):
                if "" != closing:
                    if show_tag is True:
                        line = re.sub(f'<IMG.*?alt="{closing}".*?">', f"</{tag_name}>", line, 1)
                    else:
                        line = re.sub(f'<IMG.*?alt="{closing}".*?">', '', line, 1)
        return line

    @classmethod
    def __replace_special_characters__(cls, translation_string):
        """
        Replaces "&lt;" and "&gt;" with their encoded symbols.

        :param translation_string: one segment holding the source text or the translation.
        :return: string that replaced "&lt;" and "&gt;" with "<" and ">".
        """

        tmp = re.sub("&lt;", "<", translation_string)
        tmp = re.sub("&gt;", ">", tmp)

        return tmp

    @classmethod
    def __remove_nbsp_tag__(cls, translation_string):
        """
        Removes whitespace tags within the translation.

        :param translation_string: one segment holding the source text or the translation.
        :return: segment without any whitespace tags
        """

        if 'class=field alt=&amp;nbsp; src=' in translation_string:
            tmp = re.sub('</SPAN>.*?<SPAN class=atom>', ' ', translation_string)
            return tmp
        elif '&nbsp;' in translation_string:
            tmp = re.sub('&nbsp;', '', translation_string)
            return tmp
        else:
            return translation_string

    @classmethod
    def __generate_destination_file__(cls, path):
        """
        Creates the path of the file where the processed file is to be stored. The suffix '_translated' will be added
        to the destination file automatically.

        :param path: path of the source file.
        :return: path of the destination file.
        """

        directory_name = os.path.dirname(path) + "/"
        file_name = os.path.basename(path)
        file_name = file_name.split('.')[0]
        file_name_translation = file_name + ".docx"
        path_translation = directory_name + file_name_translation

        return path_translation
